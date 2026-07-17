from flask import Flask, render_template, request, send_file
from docx import Document
from docx.table import Table
from docx.shared import Pt
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os


app = Flask(__name__)


BASE_DIR = os.path.dirname(os.path.abspath(__file__))


TEMPLATE = os.path.join(
    BASE_DIR,
    "uploads",
    "form2026.docx"
)

# =========================
# 字体设置
# =========================

def set_font(run, size, bold=True):

    run.font.name = "宋体"

    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        "宋体"
    )

    run.font.size = Pt(size)

    run.bold = bold

# =========================
# 获取全部表格
# =========================

def get_all_tables(doc):

    tables = []

    for tbl in doc._body._body.iter(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"
    ):

        tables.append(
            Table(tbl, doc._body)
        )

    return tables

# =========================
# 替换段落文字
# 保留原格式
# =========================

def replace_paragraph(paragraph, text, size):

    for run in paragraph.runs:
        run.text = ""


    run = paragraph.add_run(text)


    set_font(
        run,
        size,
        True
    )





@app.route("/")
def index():

    return render_template(
        "index.html"
    )





@app.route(
    "/generate",
    methods=["POST"]
)
def generate():


    # 客户可为空

    customer = request.form.get(
        "customer",
        ""
    )


    # 最多6个字

    customer = customer[:6]



    # 日期

    date = datetime.now().strftime(
        "%Y年%m月%d日"
    )



    # 获取菜品

    text = request.form.get(
        "dishes",
        ""
    )


    dishes = [

        x.strip()

        for x in text.splitlines()

        if x.strip()

    ]



    doc = Document(
        TEMPLATE
    )


    tables = get_all_tables(
        doc
    )

    # =========================
    # 客户 + 日期
    # 同一行
    # 时间靠右
    # 保留标题
    # =========================

    if len(tables) > 0:


        header = tables[0]


        for row in header.rows:


            for cell in row.cells:


                for p in cell.paragraphs:


                    txt = p.text.strip()



                    if (
                        "客户：" in txt
                        and
                        "2026" in txt
                    ):



                        # 清空当前行文字
                        for run in p.runs:
                            run.text = ""



                        # 设置右侧制表位

                        p.paragraph_format.tab_stops.add_tab_stop(
                            Pt(300),
                            WD_TAB_ALIGNMENT.RIGHT
                        )



                        run = p.add_run(

                            "客户："

                            +

                            customer

                            +

                            "\t"

                            +

                            date

                        )


                        set_font(
                            run,
                            14,
                            True
                        )


                        break

    # =========================
    # 查找菜品列
    # =========================

    dish_cells = []



    for table in tables[1:]:


        for row in table.rows:


            cells = row.cells


            if len(cells) >= 5:


                number = cells[0].text.strip()



                if number.isdigit():


                    n = int(number)


                    if 1 <= n <= 47:


                        # 第二列为菜品

                        dish_cells.append(
                            cells[1]
                        )

    # =========================
    # 写入菜品
    # 只改菜品
    # =========================

    for i, cell in enumerate(dish_cells):


        if i < len(dishes):


            p = cell.paragraphs[0]


            for run in p.runs:
                run.text = ""



            run = p.add_run(
                dishes[i]
            )


            set_font(
                run,
                18,
                True
            )

    # =========================
    # 保存
    # =========================

    output_dir = os.path.join(
        BASE_DIR,
        "output"
    )


    os.makedirs(
        output_dir,
        exist_ok=True
    )



    filename = (

        "蔬菜单_"

        +

        datetime.now().strftime(
            "%Y%m%d%H%M%S"
        )

        +

        ".docx"

    )


    output = os.path.join(
        output_dir,
        filename
    )



    doc.save(
        output
    )



    return send_file(
        output,
        as_attachment=True
    )

# =========================
# Railway启动
# =========================

if __name__ == "__main__":


    port = int(
        os.environ.get(
            "PORT",
            5000
        )
    )


    app.run(
        host="0.0.0.0",
        port=port
    )
