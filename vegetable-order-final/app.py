from flask import Flask, render_template, request, send_file
from docx import Document
from datetime import datetime
import os

app = Flask(__name__)
TEMPLATE = "uploads/danzi2026.docx"

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/generate", methods=["POST"])
def generate():
    customer = request.form.get("customer","")
    date = request.form.get("date","")
    dishes = [x.strip() for x in request.form.get("dishes","").splitlines() if x.strip()]

    doc = Document(TEMPLATE)

    for p in doc.paragraphs:
        if "客户：" in p.text:
            p.text = f"客户：{customer}    {date}"

    if doc.tables:
        table = doc.tables[0]
for t in doc.tables:
    first_row = [c.text for c in t.rows[0].cells]

    if "菜品" in first_row:
        table = t
        break


if table:

    for i, dish in enumerate(dishes):

        row = i + 1

        if row < len(table.rows):
            table.rows[row].cells[0].text = dish
        
        for i, dish in enumerate(dishes):
            if i + 1 < len(table.rows):
                table.rows[i+1].cells[0].text = dish

    os.makedirs("output", exist_ok=True)
    name = "蔬菜单_" + datetime.now().strftime("%Y%m%d%H%M%S") + ".docx"
    path = os.path.join("output", name)
    doc.save(path)
    return send_file(path, as_attachment=True)
    
app.config["PROPAGATE_EXCEPTIONS"] = True

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
