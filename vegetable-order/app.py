from flask import Flask, render_template, request, send_file
from docx import Document
from datetime import datetime
import os

app = Flask(__name__)
TEMPLATE = "uploads/单子2026.docx"

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/generate", methods=["POST"])
def generate():
    customer = request.form.get("customer","")
    date = request.form.get("date","")
    dishes = [x.strip() for x in request.form.get("dishes","").splitlines() if x.strip()]

    doc = Document(TEMPLATE)

    for p in doc.paragraphs:
        if "客户：" in p.text:
            p.text = f"客户：{customer}    {date}"

    if doc.tables:
        table = doc.tables[0]
        for i, dish in enumerate(dishes):
            if i + 1 < len(table.rows):
                table.rows[i+1].cells[0].text = dish

    name = "蔬菜单_" + datetime.now().strftime("%Y%m%d%H%M%S") + ".docx"
    path = os.path.join("output", name)
    doc.save(path)
    return send_file(path, as_attachment=True)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
