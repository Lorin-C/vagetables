from flask import Flask, request, send_file
from docx import Document
from docx.table import Table
from docx.shared import Pt
from docx.enum.text import WD_TAB_ALIGNMENT
from datetime import datetime
import os

app = Flask(__name__)

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE = os.path.join(BASE_DIR, "uploads", "form2026.docx")


def set_font(run, size, bold=True):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        "宋体"
    )
    run.font.size = Pt(size)
    run.bold = bold


def get_all_tables(doc):
    tables = []
    for tbl in doc._body._body.iter(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"
    ):
        tables.append(Table(tbl, doc._body))
    return tables


@app.route('/api/generate', methods=['POST'])
def generate():
    customer = request.form.get('customer', '')[:6]
    date = datetime.now().strftime('%Y年%m月%d日')
    dishes = [x.strip() for x in request.form.get('dishes','').splitlines() if x.strip()]

    doc = Document(TEMPLATE)
    tables = get_all_tables(doc)

    if tables:
        for row in tables[0].rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if '客户：' in p.text:
                        for r in p.runs:
                            r.text=''
                        p.paragraph_format.tab_stops.add_tab_stop(Pt(300), WD_TAB_ALIGNMENT.RIGHT)
                        run=p.add_run(f'客户：{customer}\t{date}')
                        set_font(run,14,True)
                        break

    cells=[]
    for table in tables[1:]:
        for row in table.rows:
            if len(row.cells)>=5 and row.cells[0].text.strip().isdigit():
                n=int(row.cells[0].text.strip())
                if 1<=n<=47:
                    cells.append(row.cells[1])

    for i,cell in enumerate(cells):
        if i<len(dishes):
            p=cell.paragraphs[0]
            for r in p.runs:
                r.text=''
            run=p.add_run(dishes[i])
            set_font(run,18,True)

    out='/tmp/蔬菜单.docx'
    doc.save(out)
    return send_file(out, as_attachment=True, download_name='蔬菜单.docx')
